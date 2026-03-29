import os
import uuid
import logging
from typing import List, Tuple
from pathlib import Path

from config import settings

logger = logging.getLogger(__name__)


def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
    """Split text into overlapping chunks by word boundaries."""
    chunk_size = chunk_size or settings.CHUNK_SIZE
    overlap = overlap or settings.CHUNK_OVERLAP

    words = text.split()
    if not words:
        return []

    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk.strip())
        if end >= len(words):
            break
        start += chunk_size - overlap

    return chunks


def extract_text_from_file(file_path: str, file_extension: str) -> str:
    """Extract raw text from various file types."""
    ext = file_extension.lower()

    if ext == ".txt" or ext == ".md":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    elif ext == ".pdf":
        return _extract_pdf(file_path)

    elif ext == ".docx":
        return _extract_docx(file_path)

    elif ext == ".csv":
        return _extract_csv(file_path)

    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_path: str) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except ImportError:
        raise ImportError("pypdf is required for PDF processing: pip install pypdf")


def _extract_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except ImportError:
        raise ImportError("python-docx is required for DOCX processing: pip install python-docx")


def _extract_csv(file_path: str) -> str:
    import csv
    rows = []
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(cell.strip() for cell in row if cell.strip()))
    return "\n".join(rows)


async def process_document(
    file_path: str,
    document_id: str,
    original_name: str,
    filename: str,
) -> Tuple[List[str], List[dict]]:
    """
    Full pipeline: extract → clean → chunk.
    Returns (chunks, chunk_metadata_list).
    """
    ext = Path(original_name).suffix.lower()
    raw_text = extract_text_from_file(file_path, ext)

    # Basic cleaning
    cleaned = "\n".join(
        line.strip() for line in raw_text.splitlines() if line.strip()
    )

    chunks = chunk_text(cleaned)
    if not chunks:
        raise ValueError("No text content could be extracted from the document.")

    chunk_metas = []
    for i, chunk in enumerate(chunks):
        chunk_metas.append({
            "chunk_id": str(uuid.uuid4()),
            "document_id": document_id,
            "filename": filename,
            "original_name": original_name,
            "chunk_index": i,
            "text": chunk,
        })

    logger.info(f"Processed '{original_name}' → {len(chunks)} chunks")
    return chunks, chunk_metas
